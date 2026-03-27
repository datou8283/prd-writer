#!/usr/bin/env python3
"""Generate the Homepage Feature Navigation PRD as .docx with embedded flow PNGs."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DOCS = "/Users/dannychen/Documents/aster-docs/docs"

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Helvetica Neue"
font.size = Pt(10)

# ============================================================
# §1 Feature Overview
# ============================================================
doc.add_heading("Homepage Feature Navigation Menu", level=0)

doc.add_heading("1. Feature Overview", level=1)
doc.add_paragraph(
    "Redesign the Aster mobile app's hamburger menu from a simple settings drawer into a full-screen "
    "feature navigation hub. The new navigation page provides a categorized icon grid covering all "
    "platform features (Deposit, Withdraw, Perpetual, Spot, Shield, Earn, Staking, Explorer, Referral, "
    "Rocket Launch, Trade & Earn, Points, Airdrop, Leaderboard, and more), user-customizable shortcut "
    "ordering (persisted on device), a settings sub-page for relocated menu items, and WebView integration "
    "for features that exist on web but not yet natively in the app. This bridges the feature gap between "
    "Aster's web and mobile platforms, enabling all users to access every platform feature from the app."
)

# ============================================================
# §2 Background & Motivation
# ============================================================
doc.add_heading("2. Background & Motivation", level=1)

doc.add_paragraph(
    "Aster Exchange currently has a significant feature gap between its web and mobile platforms. "
    "Several features — including Staking, Explorer, Leaderboard, Points, Airdrop, and campaign-related "
    "functions — are available on web but entirely absent from the mobile app. Users who want to access "
    "these features must leave the app and use a mobile browser, creating friction and reducing engagement."
)
doc.add_paragraph(
    "The current hamburger menu only contains settings-related items (Security, Language, Help, Referral, "
    "Give Feedback, Privacy Policy, Terms & Conditions, About) and provides no feature discovery or "
    "navigation capability. This is out of step with industry standards."
)
doc.add_paragraph(
    "Per the competitor analysis (see homepage-navigation-competitor-analysis.docx), all major exchanges "
    "(Binance, OKX, Bybit) have moved to full-screen profile/account pages with categorized feature grids "
    "and user-customizable shortcuts. OKX's profile page model — with separated Shortcuts, Manage Assets, "
    "and Trade sections — is the closest match to Aster's needs and has been selected as the reference design, "
    "adapted for Aster's DEX context (wallet address replaces username, no KYC badges)."
)

# ============================================================
# §3 Scope
# ============================================================
doc.add_heading("3. Scope", level=1)

doc.add_heading("In Scope (v1)", level=2)
in_scope = [
    "Full-screen feature navigation page replacing the current hamburger menu drawer",
    "Categorized feature icon grid with section headers (Manage Assets, Trade, Earn, Explore, Rewards, Support)",
    "Feature list, categories, and routing hardcoded in the app bundle (updated with each app release)",
    "User-customizable shortcut ordering with drag-to-reorder edit mode, persisted to device local storage (AsyncStorage / MMKV)",
    "Settings sub-page relocated from hamburger menu (Security, Language, Help, Privacy Policy, T&C, About)",
    "WebView container for features that exist on web but not natively in app, with auth token injection",
    "i18n for all feature names, category labels, and error messages (EN, zh-Hans, zh-Hant)",
    "Referral and Give Feedback moved from settings to feature navigation grid",
]
for item in in_scope:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Out of Scope (v2+)", level=2)
out_scope = [
    ("Server-driven feature configuration", "v1 hardcodes feature list in app; v2 adds config API so admin can update features, ordering, and visibility without app release"),
    ("Feature badges (NEW, HOT) on icon grid items", "requires server-driven config to control badge state dynamically; add in v2 alongside config API"),
    ("Navigation config API (public) and user shortcuts API (authenticated)", "v1 stores all data client-side; v2 adds server sync for cross-device shortcut persistence and remote feature config"),
    ("Bottom navigation bar changes", "independent PRD — current 5-tab layout stays"),
    ("New feature implementations", "this PRD covers navigation entry points only, not the features themselves"),
    ("ML/AI-powered feature recommendations", "Binance-style AI personalization is v2+"),
    ("Hamburger icon → profile avatar", "product decision to keep hamburger icon for v1"),
    ("Homepage widget customization", "Bybit-style widget layout is over-engineered for v1; simple shortcut reorder suffices"),
]
for title, reason in out_scope:
    doc.add_paragraph(f"{title} — {reason}", style="List Bullet")

# ============================================================
# §4 User Stories
# ============================================================
doc.add_heading("4. User Stories", level=1)

stories = [
    ["Trader", "Open the hamburger menu from the homepage", "See a full-screen feature navigation page with all platform features organized by category"],
    ["Trader", "Tap a feature icon for an app-native feature (e.g., Perpetual)", "Navigate directly to the Perpetual trading screen"],
    ["Trader", "Tap a feature icon for a web-only feature (e.g., Points)", "Feature opens in an in-app WebView with auth token injected — no re-login needed"],
    ["Trader", "Tap the edit button on the Shortcuts section", "Enter drag-to-reorder mode to customize shortcut arrangement"],
    ["Trader", "Drag shortcuts to reorder and tap Done", "New shortcut order is saved to device and persists across app restarts"],
    ["Trader", "Tap the settings gear icon on the navigation page", "Navigate to the settings sub-page with Security, Language, Help, Privacy Policy, T&C, About"],
    ["Trader", "Reinstall the app or switch devices", "Shortcut customization is reset to default (device-local storage only in v1)"],
]

table = doc.add_table(rows=len(stories) + 1, cols=3)
table.style = "Light Grid Accent 1"
header = table.rows[0]
for j, h in enumerate(["Role", "Action", "Outcome"]):
    header.cells[j].text = h
    for run in header.cells[j].paragraphs[0].runs:
        run.bold = True

for i, row_data in enumerate(stories):
    for j, cell_text in enumerate(row_data):
        table.rows[i + 1].cells[j].text = cell_text

# ============================================================
# §5 User Flows
# ============================================================
doc.add_heading("5. User Flows", level=1)

doc.add_heading("Flow 1: Menu Navigation to Feature", level=2)
doc.add_paragraph(
    "This flow covers the primary user journey from the homepage to any platform feature via the "
    "new navigation page. The decision point routes the user to either a native app screen or a "
    "WebView for web-only features."
)
doc.add_picture(os.path.join(DOCS, "flow-menu-navigation.png"), width=Inches(5.5))

doc.add_heading("Flow 2: Customize Shortcuts", level=2)
doc.add_paragraph(
    "This flow covers how a user enters edit mode, rearranges shortcuts via drag-and-drop, "
    "and saves the new arrangement to device local storage."
)
doc.add_picture(os.path.join(DOCS, "flow-customize-shortcuts.png"), width=Inches(5.0))

doc.add_heading("Flow 3: Settings Access", level=2)
doc.add_paragraph(
    "This flow shows how a user accesses the settings sub-page, which now lives inside the "
    "navigation page rather than being the primary content of the hamburger menu."
)
doc.add_picture(os.path.join(DOCS, "flow-settings-access.png"), width=Inches(4.0))

# ============================================================
# §6 Functional Requirements
# ============================================================
doc.add_heading("6. Functional Requirements", level=1)

# 6.1
doc.add_heading("6.1 Navigation Page UI", level=2)
reqs = [
    "The hamburger icon (top-left of homepage header) must open the navigation page via push navigation (slide-in from right, not drawer/overlay).",
    "The navigation page must be a full-screen page on the navigation stack with a back button.",
    "Header section must display: wallet address (truncated, e.g., 0x6b...2EDf) on the left, settings gear icon on the right.",
    "Below the header: scrollable content area with categorized feature grid sections.",
    "Each section has a section header label (e.g., 'Manage Assets', 'Trade') and a 4-column icon grid below it.",
    "The Shortcuts section (topmost) must have an edit (pencil) icon in the section header to enter customization mode.",
    "Page background must match the app's dark theme.",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

# 6.2
doc.add_heading("6.2 Feature Icon Grid", level=2)
reqs = [
    "Each feature icon must be displayed as: icon image (bundled asset) + text label below.",
    "Grid layout: 4 columns, variable rows per category. Icon size: 48×48dp, label below.",
    "Feature labels must be i18n — displayed in the user's selected language (EN, zh-Hans, zh-Hant).",
    "Tap behavior depends on route_type (hardcoded per feature in the app bundle):",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

# route_type table
rt_table = doc.add_table(rows=4, cols=3)
rt_table.style = "Light Grid Accent 1"
rt_headers = ["route_type", "Behavior", "Example"]
for j, h in enumerate(rt_headers):
    rt_table.rows[0].cells[j].text = h
    for run in rt_table.rows[0].cells[j].paragraphs[0].runs:
        run.bold = True

rt_data = [
    ["native", "Deep link to in-app screen via route_target path", "Perpetual → /perpetual/BTCUSDT"],
    ["webview", "Open WebView with route_target URL, inject auth token", "Points → https://app.aster.exchange/points"],
    ["disabled", "Show grayed-out icon, non-tappable. Optional tooltip: 'Coming soon'", "N/A — reserve for future features"],
]
for i, row in enumerate(rt_data):
    for j, cell in enumerate(row):
        rt_table.rows[i + 1].cells[j].text = cell

# 6.3
doc.add_heading("6.3 Feature Categories and List", level=2)
doc.add_paragraph(
    "Features are organized into the following categories. In v1, the category structure, feature list, "
    "and routing configuration are all hardcoded in the app bundle and updated with each app release."
)

cat_table = doc.add_table(rows=7, cols=3)
cat_table.style = "Light Grid Accent 1"
cat_headers = ["Category", "Features (v1)", "Notes"]
for j, h in enumerate(cat_headers):
    cat_table.rows[0].cells[j].text = h
    for run in cat_table.rows[0].cells[j].paragraphs[0].runs:
        run.bold = True

cat_data = [
    ["Shortcuts", "User-customizable — default: Deposit, Withdraw, Perpetual, Spot", "Only section with user customization. Default set hardcoded."],
    ["Manage Assets", "Deposit, Withdraw", "Core wallet operations"],
    ["Trade", "Perpetual, Spot, Shield", "Trading entry points — deep link to respective trading screens"],
    ["Earn & Rewards", "Earn, Staking, Trade & Earn, Points, Airdrop, Rocket Launch", "Passive income products. Campaign and incentive features"],
    ["Explore", "Explorer, Referral, Leaderboard", "Discovery and community features. Referral relocated from old menu."],
    ["Support", "Give Feedback", "Relocated from old menu. May expand with Help Center link."],
]
for i, row in enumerate(cat_data):
    for j, cell in enumerate(row):
        cat_table.rows[i + 1].cells[j].text = cell

# Feature routing table
doc.add_paragraph("\nFeature routing configuration (hardcoded in v1):")

fr_table = doc.add_table(rows=16, cols=3)
fr_table.style = "Light Grid Accent 1"
fr_headers = ["Feature", "route_type", "route_target"]
for j, h in enumerate(fr_headers):
    fr_table.rows[0].cells[j].text = h
    for run in fr_table.rows[0].cells[j].paragraphs[0].runs:
        run.bold = True

fr_data = [
    ["Deposit", "native", "/wallet/deposit"],
    ["Withdraw", "native", "/wallet/withdraw"],
    ["Perpetual", "native", "/perpetual/BTCUSDT"],
    ["Spot", "native", "/spot/BTCUSDT"],
    ["Shield", "native", "/shield"],
    ["Earn", "webview", "https://app.aster.exchange/earn"],
    ["Staking", "webview", "https://app.aster.exchange/staking"],
    ["Trade & Earn", "webview", "https://app.aster.exchange/trade-and-earn"],
    ["Points", "webview", "https://app.aster.exchange/points"],
    ["Airdrop", "webview", "https://app.aster.exchange/airdrop"],
    ["Rocket Launch", "webview", "https://app.aster.exchange/rocket-launch"],
    ["Explorer", "webview", "https://app.aster.exchange/explorer"],
    ["Referral", "webview", "https://app.aster.exchange/referral"],
    ["Leaderboard", "webview", "https://app.aster.exchange/leaderboard"],
    ["Give Feedback", "webview", "https://app.aster.exchange/feedback"],
]
for i, row in enumerate(fr_data):
    for j, cell in enumerate(row):
        fr_table.rows[i + 1].cells[j].text = cell

# 6.4
doc.add_heading("6.4 User-Customizable Shortcuts", level=2)
reqs = [
    "The Shortcuts section must support user customization of which features appear and their order.",
    "Edit mode is triggered by tapping the pencil icon in the Shortcuts section header.",
    "In edit mode: all available features are shown in a grid. Currently selected shortcuts have a checkmark; unselected are dimmed.",
    "User can tap to add/remove features from shortcuts, and long-press + drag to reorder.",
    "Maximum shortcuts: 8 (2 rows of 4). Minimum: 4 (1 row).",
    "A 'Reset to default' button must be available in edit mode to restore the hardcoded default set.",
    "Tapping 'Done' saves the new arrangement to device local storage (AsyncStorage or MMKV).",
    "The shortcut arrangement persists across app restarts but is device-local only (not synced across devices in v1).",
    "Default shortcut set for new users / fresh installs: Deposit, Withdraw, Perpetual, Spot.",
    "On app reinstall or clear data: shortcuts reset to default.",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

# 6.5
doc.add_heading("6.5 Settings Sub-Page", level=2)
reqs = [
    "Accessible via the gear icon in the navigation page header.",
    "Push navigation to a standard list-style settings page.",
    "Items (all relocated from the current hamburger menu): Security, Language, Help, Privacy Policy, Terms & Conditions, About.",
    "Each item navigates to its existing detail screen — no changes to the detail screens themselves.",
    "The old hamburger menu drawer is fully replaced and should be removed from the codebase.",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

# 6.6
doc.add_heading("6.6 WebView Container", level=2)
reqs = [
    "Full-screen WebView must be used for features with route_type='webview'.",
    "WebView page layout: top navigation bar (back button + feature name as title) + WebView content below.",
    "Auth token must be injected automatically so the user does not need to re-login in the WebView.",
    "Auth injection method: set JWT token as a cookie on the WebView domain before loading the URL. Fallback: inject via JavaScript postMessage after page load.",
    "WebView must support navigation within the web app (links within the same domain open in the WebView, external links open in system browser).",
    "On load failure (network error, timeout): show an error page with a 'Retry' button.",
    "WebView must support deep link callbacks from web → app (e.g., web page triggers an aster:// deep link to navigate back to a native screen).",
    "Loading indicator must be shown while WebView content loads.",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

# ============================================================
# §7 API Specification
# ============================================================
doc.add_heading("7. API Specification", level=1)

doc.add_paragraph(
    "v1 does not require any new backend API endpoints. The feature list, categories, and routing "
    "configuration are hardcoded in the app bundle. User shortcut customization is persisted to "
    "device local storage (AsyncStorage / MMKV)."
)

doc.add_heading("v2 API Preview (out of scope for v1)", level=2)
doc.add_paragraph(
    "The following endpoints are planned for v2 to enable server-driven configuration and cross-device "
    "shortcut synchronization. They are documented here for architectural awareness only."
)

v2_apis = [
    ["Endpoint", "Method", "Description"],
    ["GET /api/v1/app/navigation-config", "Public", "Retrieve server-driven feature list, categories, ordering, visibility, and badges. Replaces hardcoded config."],
    ["GET /api/v1/user/navigation-shortcuts", "Authenticated", "Retrieve user's customized shortcut arrangement from server (cross-device sync)."],
    ["PUT /api/v1/user/navigation-shortcuts", "Authenticated", "Save user's shortcut arrangement to server. Replaces device-local storage."],
]

t = doc.add_table(rows=len(v2_apis), cols=3)
t.style = "Light Grid Accent 1"
for i, row in enumerate(v2_apis):
    for j, cell in enumerate(row):
        t.rows[i].cells[j].text = cell
        if i == 0:
            for run in t.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph(
    "\nMigration path: when v2 APIs ship, the app should check for existing local shortcuts and "
    "upload them to the server on first launch after upgrade, then switch to server as the source of truth."
)

# ============================================================
# §8 Error Handling
# ============================================================
doc.add_heading("8. Error Handling", level=1)

doc.add_paragraph(
    "With no backend APIs in v1, error handling is focused on the client side: WebView load failures, "
    "local storage issues, and disabled feature handling."
)

err_data = [
    ["Code", "EN", "zh-Hans", "zh-Hant", "Trigger"],
    ["50006", "Feature not available", "功能暂不可用", "功能暫不可用", "Tap disabled feature icon"],
    ["50007", "WebView load failed", "页面加载失败", "頁面載入失敗", "WebView URL fails to load (network/timeout)"],
    ["50009", "Local save failed", "保存失败", "儲存失敗", "AsyncStorage/MMKV write failure (extremely rare)"],
]

err_table = doc.add_table(rows=len(err_data), cols=5)
err_table.style = "Light Grid Accent 1"
for i, row in enumerate(err_data):
    for j, cell in enumerate(row):
        err_table.rows[i].cells[j].text = cell
        if i == 0:
            for run in err_table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph("")
doc.add_heading("Edge cases", level=2)
edges = [
    "User opens a WebView feature while offline: show error page with 'Retry' button. When connectivity resumes, user taps Retry to reload.",
    "Local storage corrupted or unreadable: fall back to hardcoded default shortcuts silently. No error UI shown — user simply sees the default set.",
    "App update adds new features to the hardcoded list: new features appear in the grid automatically. User's custom shortcuts are preserved; new features are only in the category grid, not auto-added to shortcuts.",
    "Auth token expires during WebView session: WebView should detect 401 responses and prompt user to re-authenticate (redirect to wallet connection flow).",
]
for e in edges:
    doc.add_paragraph(e, style="List Bullet")

# ============================================================
# §9 Non-Functional Requirements
# ============================================================
doc.add_heading("9. Non-Functional Requirements", level=1)

nfr_data = [
    ["Dimension", "Requirement"],
    ["Performance — navigation page render", "Time to interactive ≤ 300ms from hamburger icon tap. No async data loading needed (config is bundled)."],
    ["Performance — drag animation", "60fps during drag-to-reorder. Use React Native Reanimated or equivalent for gesture handling."],
    ["Performance — local storage", "AsyncStorage/MMKV read/write ≤ 10ms. Shortcut data is small (< 1KB)."],
    ["i18n", "All feature names, category labels, and error messages must be provided in EN, zh-Hans, zh-Hant. No hardcoded user-facing strings."],
    ["Accessibility", "All feature icons must have proper accessibility labels for screen readers. Edit mode drag handles must be accessible via VoiceOver/TalkBack."],
    ["Security — WebView", "Auth token must not be exposed in URL query parameters. Use cookie injection (HttpOnly, Secure, SameSite=Strict) or JavaScript postMessage."],
    ["Image assets", "Feature icons bundled in the app (not fetched from CDN in v1). Icon format: PNG, 96×96px @3x."],
    ["Offline support", "Navigation page renders entirely from local data — no network required. WebView features require connectivity."],
    ["Storage", "Shortcut customization stored in AsyncStorage/MMKV under a namespaced key (e.g., @aster/navigation_shortcuts). Device-local only in v1."],
]

nfr_table = doc.add_table(rows=len(nfr_data), cols=2)
nfr_table.style = "Light Grid Accent 1"
for i, row in enumerate(nfr_data):
    for j, cell in enumerate(row):
        nfr_table.rows[i].cells[j].text = cell
        if i == 0:
            for run in nfr_table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

# ============================================================
# §10 Open Questions
# ============================================================
doc.add_heading("10. Open Questions", level=1)

questions = [
    ["#", "Question", "Impact", "Proposed default"],
    ["1", "Is the feature list complete? The original request said 'etc.' — are there additional features beyond the 15 listed?", "Affects category structure and grid layout", "Proceed with 15 listed features; add more in next app release"],
    ["2", "For each feature, what is the exact route_type (native vs webview)? Which features currently have native app screens?", "Determines WebView scope and auth token injection priority", "Perpetual, Spot, Deposit, Withdraw = native; all others = webview for v1 (see §6.3 table)"],
    ["3", "WebView auth token injection method: cookie injection vs JavaScript postMessage vs URL fragment?", "Security architecture and web team coordination", "Cookie injection (HttpOnly, Secure) — most secure, requires web app to read auth from cookie"],
    ["4", "Should the Shortcuts section be above or below the wallet address / profile header?", "UI hierarchy and visual weight", "Below header, above all other categories (matches OKX pattern)"],
    ["5", "Which local storage library: AsyncStorage vs MMKV?", "Performance and migration complexity", "MMKV (faster synchronous reads, better for UI-critical data). Fall back to AsyncStorage if MMKV not already in the project."],
]

q_table = doc.add_table(rows=len(questions), cols=4)
q_table.style = "Light Grid Accent 1"
for i, row in enumerate(questions):
    for j, cell in enumerate(row):
        q_table.rows[i].cells[j].text = cell
        if i == 0:
            for run in q_table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

# Save
output_path = os.path.join(DOCS, "homepage-navigation-prd.docx")
doc.save(output_path)
print(f"Saved to {output_path}")
