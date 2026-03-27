#!/usr/bin/env python3
"""Generate PostHog event tracking spec for Homepage Navigation feature."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Helvetica Neue"
font.size = Pt(10)

doc.add_heading("PostHog Event Tracking Spec: Homepage Feature Navigation", level=0)
doc.add_paragraph(
    "Feature: Homepage Feature Navigation Menu\n"
    "PRD: homepage-navigation-prd.docx\n"
    "Date: 2026-03-27\n"
    "Product prefix: navigation_"
)

doc.add_heading("Events Overview", level=1)
doc.add_paragraph(
    "8 events covering page views, core actions, configuration choices, and error tracking. "
    "v1 has no backend APIs — all config is bundled, shortcuts are device-local."
)

# Helper to add event tables
def add_event(event_name, trigger, properties, business_goal=""):
    doc.add_heading(event_name, level=2)

    # Header row + property rows
    rows = len(properties) + 1  # header + data
    if not properties:
        rows = 2  # header + single dash row

    t = doc.add_table(rows=rows, cols=6)
    t.style = "Light Grid Accent 1"

    headers = ["Event name", "Trigger point", "Property", "Type", "Property value", "Business goal"]
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for run in t.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    if not properties:
        t.rows[1].cells[0].text = event_name
        t.rows[1].cells[1].text = trigger
        t.rows[1].cells[2].text = "-"
        t.rows[1].cells[3].text = "-"
        t.rows[1].cells[4].text = "-"
        t.rows[1].cells[5].text = business_goal
    else:
        for i, prop in enumerate(properties):
            row = t.rows[i + 1]
            if i == 0:
                row.cells[0].text = event_name
                row.cells[1].text = trigger
                row.cells[5].text = business_goal
            else:
                row.cells[0].text = ""
                row.cells[1].text = ""
                row.cells[5].text = ""
            row.cells[2].text = prop[0]  # property name
            row.cells[3].text = prop[1]  # type
            row.cells[4].text = prop[2]  # value

    doc.add_paragraph("")  # spacer


# ============================================================
# Event 1: navigation_menu_viewed
# ============================================================
add_event(
    "navigation_menu_viewed",
    "When the full-screen navigation page finishes rendering after user taps the hamburger icon",
    [
        ("source", "Enum", "homepage_hamburger"),
        ("feature_count", "Num", "Total number of features in the grid"),
        ("shortcut_count", "Num", "Number of shortcuts displayed"),
        ("is_customized", "Boolean", "true if user has custom shortcuts in local storage, false if using defaults"),
    ],
    "Track navigation page engagement and adoption."
)

# ============================================================
# Event 2: navigation_feature_clicked
# ============================================================
add_event(
    "navigation_feature_clicked",
    "When user taps a feature icon in the navigation page grid (before navigation occurs)",
    [
        ("feature_id", "String", "Feature identifier, e.g., 'perpetual', 'deposit', 'points'"),
        ("feature_name", "String", "Feature display name in user's locale"),
        ("category_id", "String", "Category of the feature, e.g., 'trade', 'earn', 'explore'"),
        ("route_type", "Enum", "native | webview | disabled"),
        ("position_index", "Num", "0-based index of the feature within its category grid"),
        ("is_shortcut", "Boolean", "true if tapped from the Shortcuts section, false if from a category section"),
    ],
    "Feature usage heatmap. Identify most/least used features."
)

# ============================================================
# Event 3: navigation_edit_mode_entered
# ============================================================
add_event(
    "navigation_edit_mode_entered",
    "When user taps the pencil/edit icon on the Shortcuts section header",
    [
        ("current_shortcut_count", "Num", "Number of shortcuts before editing"),
        ("current_shortcut_ids", "JSON String", 'Ordered list of current shortcut feature_ids, e.g., ["deposit","withdraw","perpetual","spot"]'),
    ],
    "Measure customization adoption. Track which default sets users want to change."
)

# ============================================================
# Event 4: navigation_shortcuts_saved
# ============================================================
add_event(
    "navigation_shortcuts_saved",
    "When user taps Done and shortcuts are saved to device local storage",
    [
        ("shortcut_count", "Num", "Number of shortcuts in the new arrangement"),
        ("shortcut_ids", "JSON String", 'Ordered list of new shortcut feature_ids'),
        ("added_ids", "JSON String", "Feature IDs added in this edit session"),
        ("removed_ids", "JSON String", "Feature IDs removed in this edit session"),
        ("reordered", "Boolean", "true if only order changed (no add/remove), false otherwise"),
    ],
    "Track what users customize. Identify popular shortcuts to inform default set."
)

# ============================================================
# Event 5: navigation_shortcuts_reset
# ============================================================
add_event(
    "navigation_shortcuts_reset",
    "When user taps 'Reset to default' button in edit mode and confirms",
    [
        ("previous_shortcut_ids", "JSON String", "Shortcut IDs before reset"),
    ],
    "Track reset frequency — high reset rate suggests poor defaults or confusing edit UX."
)

# ============================================================
# Event 6: navigation_settings_viewed
# ============================================================
add_event(
    "navigation_settings_viewed",
    "When the settings sub-page finishes rendering after user taps the gear icon",
    [],
    "Track settings page access to ensure the relocated settings remain discoverable."
)

# ============================================================
# Event 7: navigation_webview_opened
# ============================================================
add_event(
    "navigation_webview_opened",
    "When the WebView container finishes loading the target URL (onLoad event)",
    [
        ("feature_id", "String", "Feature that triggered the WebView"),
        ("url", "String", "WebView target URL"),
        ("load_duration_ms", "Num", "Time from tap to WebView onLoad event"),
        ("auth_injected", "Boolean", "true if auth token was successfully injected"),
    ],
    "Monitor WebView performance and auth injection success rate."
)

# ============================================================
# Event 8: navigation_webview_error
# ============================================================
add_event(
    "navigation_webview_error",
    "When the WebView fails to load (onError event or HTTP error status)",
    [
        ("feature_id", "String", "Feature that triggered the WebView"),
        ("url", "String", "WebView target URL"),
        ("error_type", "Enum", "network_error | timeout | http_error | auth_failure"),
        ("http_status", "Num", "HTTP status code if applicable, 0 otherwise"),
        ("retry_count", "Num", "Number of retry attempts before this error event"),
    ],
    "Monitor WebView reliability. Alert on high error rate for specific features."
)

# Save
output_path = "/Users/dannychen/Documents/aster-docs/docs/homepage-navigation-posthog-events.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
