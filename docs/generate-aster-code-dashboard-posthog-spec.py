#!/usr/bin/env python3
"""Generate PostHog event tracking spec for Aster Code Dashboard feature."""

from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Helvetica Neue"
font.size = Pt(10)

doc.add_heading("PostHog Event Tracking Spec: Aster Code Dashboard", level=0)
doc.add_paragraph(
    "Feature: Aster Code Dashboard\n"
    "PRD: aster-code-dashboard-prd.docx\n"
    "Date: 2026-04-07\n"
    "Product prefix: aster_code_"
)

doc.add_heading("Events Overview", level=1)
doc.add_paragraph(
    "8 events covering page views, chart interactions, table interactions, and error tracking. "
    "All events fire on the public dashboard — no authentication required."
)


def add_event(event_name, trigger, properties, business_goal=""):
    doc.add_heading(event_name, level=2)

    rows = len(properties) + 1
    if not properties:
        rows = 2

    t = doc.add_table(rows=rows, cols=6)
    t.style = "Light Grid Accent 1"

    headers = ["Event name", "Trigger point", "Property", "Type", "Property value", "Business goal"]
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for run in t.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    if not properties:
        t.rows[1].cells[0].text = event_name
        t.rows[1].cells[1].text = trigger
        t.rows[1].cells[2].text = "-"
        t.rows[1].cells[3].text = "-"
        t.rows[1].cells[4].text = "-"
        t.rows[1].cells[5].text = business_goal
    else:
        for i, prop in enumerate(properties):
            row = t.rows[i + 1]
            if i == 0:
                row.cells[0].text = event_name
                row.cells[1].text = trigger
                row.cells[5].text = business_goal
            else:
                row.cells[0].text = ""
                row.cells[1].text = ""
                row.cells[5].text = ""
            row.cells[2].text = prop[0]
            row.cells[3].text = prop[1]
            row.cells[4].text = prop[2]

    doc.add_paragraph("")


# ============================================================
# Event 1: aster_code_dashboard_viewed
# ============================================================
add_event(
    "aster_code_dashboard_viewed",
    "When the Aster Code Dashboard page finishes initial data load and renders KPI cards",
    [
        ("load_duration_ms", "Num", "Time from navigation start to first contentful paint"),
        ("total_codes", "Num", "Total number of Aster Codes returned by overview API"),
        ("data_last_updated", "String", "ISO 8601 timestamp of last pipeline run from API response"),
        ("referrer", "String", "Document referrer URL (where the user came from), empty string if direct"),
        ("utm_source", "String", "UTM source parameter if present, empty string otherwise"),
        ("utm_medium", "String", "UTM medium parameter if present, empty string otherwise"),
    ],
    "Track dashboard adoption, traffic sources, and data freshness perception."
)

# ============================================================
# Event 2: aster_code_chart_switched
# ============================================================
add_event(
    "aster_code_chart_switched",
    "When user clicks a time range button (W/M/Q/D/ALL) or chart type toggle (line/bar/area) on any chart",
    [
        ("chart_name", "String", "Chart identifier, e.g., 'volume_overall', 'fees_by_code', 'pie_chart', 'avg_fee_per_user'"),
        ("switch_type", "Enum", "time_range | chart_type"),
        ("from_value", "String", "Previous selection, e.g., 'M' or 'bar'"),
        ("to_value", "String", "New selection, e.g., 'W' or 'line'"),
    ],
    "Understand preferred time horizons and chart formats. Inform default selections."
)

# ============================================================
# Event 3: aster_code_pie_dimension_switched
# ============================================================
add_event(
    "aster_code_pie_dimension_switched",
    "When user clicks a dimension tab (Fees/Volume/Trades/Users) on the pie chart",
    [
        ("from_dimension", "Enum", "fees | volume | trades | users"),
        ("to_dimension", "Enum", "fees | volume | trades | users"),
    ],
    "Track which metric dimensions users care about most for code comparison."
)

# ============================================================
# Event 4: aster_code_legend_toggled
# ============================================================
add_event(
    "aster_code_legend_toggled",
    "When user clicks a code name in a chart legend to show/hide it, or clicks 'Deselect all'",
    [
        ("chart_name", "String", "Chart identifier where legend was toggled"),
        ("action", "Enum", "show | hide | deselect_all"),
        ("code_name", "String", "Aster Code name toggled (empty string for deselect_all)"),
        ("visible_code_count", "Num", "Number of codes visible after this action"),
    ],
    "Identify which codes users focus on. High deselect_all usage suggests too many codes — validate top-N approach."
)

# ============================================================
# Event 5: aster_code_table_sorted
# ============================================================
add_event(
    "aster_code_table_sorted",
    "When user clicks a column header in the Aster Code Stats table to change sort",
    [
        ("sort_by", "Enum", "code | volume | fees | trade_count_30d | total_users"),
        ("sort_order", "Enum", "asc | desc"),
        ("previous_sort_by", "String", "Previous sort column"),
    ],
    "Track which metrics users rank by. Informs default sort order."
)

# ============================================================
# Event 6: aster_code_table_paginated
# ============================================================
add_event(
    "aster_code_table_paginated",
    "When user clicks a pagination button (previous/next/page number) in the stats table",
    [
        ("from_page", "Num", "Page number before navigation"),
        ("to_page", "Num", "Page number after navigation"),
        ("total_pages", "Num", "Total number of pages available"),
        ("sort_by", "String", "Current sort column when paginating"),
    ],
    "Track how deep users browse. Low page 2+ visits may suggest table is too long or users only care about top codes."
)

# ============================================================
# Event 7: aster_code_address_copied
# ============================================================
add_event(
    "aster_code_address_copied",
    "When user clicks the copy icon next to a wallet address in the stats table",
    [
        ("code_name", "String", "Aster Code name whose address was copied"),
        ("address", "String", "Full wallet address copied"),
        ("table_row_index", "Num", "0-based row index on the current page"),
        ("page", "Num", "Current page number"),
    ],
    "Track on-chain verification behavior. High copy rate indicates users verify addresses independently."
)

# ============================================================
# Event 8: aster_code_api_error
# ============================================================
add_event(
    "aster_code_api_error",
    "When any dashboard API call returns a non-zero error code or network failure",
    [
        ("endpoint", "String", "API endpoint path, e.g., '/api/v1/aster-code/series'"),
        ("error_code", "Num", "Error code from API response (60001-60005), or 0 for network error"),
        ("http_status", "Num", "HTTP status code (0 for network failure)"),
        ("error_type", "Enum", "api_error | network_error | timeout"),
        ("retry_count", "Num", "Number of retry attempts before this error event fires"),
        ("chart_name", "String", "Chart or section that triggered the failed request (empty if overview)"),
    ],
    "Monitor API reliability. Alert on elevated error rates. Track retry success rate."
)

# ============================================================
# Summary
# ============================================================
doc.add_heading("Event Summary", level=1)

summary = [
    ["Event", "Trigger Type", "Category"],
    ["aster_code_dashboard_viewed", "System (page load)", "Page view"],
    ["aster_code_chart_switched", "User action (click)", "Configuration"],
    ["aster_code_pie_dimension_switched", "User action (click)", "Configuration"],
    ["aster_code_legend_toggled", "User action (click)", "Configuration"],
    ["aster_code_table_sorted", "User action (click)", "Core action"],
    ["aster_code_table_paginated", "User action (click)", "Core action"],
    ["aster_code_address_copied", "User action (click)", "Core action"],
    ["aster_code_api_error", "System (API response)", "Error tracking"],
]

t = doc.add_table(rows=len(summary), cols=3)
t.style = "Light Grid Accent 1"
for i, row in enumerate(summary):
    for j, cell in enumerate(row):
        t.rows[i].cells[j].text = cell
        if i == 0:
            for run in t.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph(
    "\nTotal: 8 events (1 page view, 4 configuration/interaction, 2 core actions, 1 error tracking)."
)

# Save
output_path = "/Users/dannychen/Documents/aster-docs/docs/aster-code-dashboard-posthog-events.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
