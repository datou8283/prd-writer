#!/usr/bin/env python3
"""Generate the homepage navigation competitor analysis .docx report."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
font = style.font
font.name = "Helvetica Neue"
font.size = Pt(10)

# Title
title = doc.add_heading("Competitor Analysis: Mobile App Homepage Navigation", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph(
    "Feature: Homepage Feature Navigation Menu Redesign\n"
    "Date: 2026-03-27\n"
    "Competitors analyzed: Binance, OKX, Bybit, Hyperliquid"
)

# ============================================================
# 1. Executive Summary
# ============================================================
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This report analyzes how four major exchanges (Binance, OKX, Bybit, Hyperliquid) "
    "structure their mobile app homepage navigation and feature discovery. "
    "The key finding is that all major CEXs have moved from simple hamburger menus to "
    "rich profile/account pages with categorized feature grids and customizable shortcuts. "
    "Recommendation: Aster should adopt OKX's profile-page-with-shortcuts model adapted "
    "for a DEX context (wallet address replaces username, no KYC badges), with server-driven "
    "feature configuration and v1 user customization."
)

# ============================================================
# 2. Feature Comparison Matrix
# ============================================================
doc.add_heading("2. Feature Comparison Matrix", level=1)

matrix_data = [
    ["Dimension", "Binance", "OKX", "Bybit", "Hyperliquid"],
    ["Menu trigger", "Profile icon (top-left)", "Profile avatar (top-left)", "Profile icon (top-left)", "No dedicated menu (sidebar nav on web)"],
    ["Menu format", "Full-screen Account Center", "Full-screen profile page (push nav)", "Full-screen My Page", "N/A — web-only, minimal nav bar"],
    ["Feature grid", "Yes — shortcuts + recommended features", "Yes — Shortcuts + Manage Assets + Trade sections", "Yes — customizable homepage widgets", "No — flat nav bar (Trade, Portfolio)"],
    ["Feature categories", "Shortcuts, Recommended, Services", "Shortcuts, Manage Assets, Trade", "Varies by Lite/Pro mode", "Trade, Portfolio only"],
    ["Shortcuts customizable", "Yes — manage favorite mini apps", "Yes — edit button on Shortcuts section", "Yes — customizable homepage (Pro mode)", "No"],
    ["Shortcut count", "~8 default shortcuts", "4 default shortcuts (expandable)", "Varies by mode", "N/A"],
    ["Settings location", "Settings icon in Account Center", "Profile and settings sub-page", "My Page → Settings", "Gear icon on web"],
    ["Lite/Pro modes", "Yes — switch in Account Center", "No (single mode)", "Yes — toggle at top of app", "No"],
    ["WebView for web-only features", "Yes (some Earn products)", "Yes (DeFi, Web3 features)", "Yes (Web3 section)", "N/A — all features are web"],
    ["Bottom tab count", "5 (Home, Markets, Trade, Futures, Wallets)", "5 (Home, Discover, Trade, Web3, Wallet)", "5 (Home, Markets, Trade, Copy, Assets)", "3-4 (Trade, Portfolio, more)"],
    ["Taps to reach any feature", "2-3 (homepage → Account Center → feature)", "2-3 (homepage → profile → feature)", "2-3 (homepage → menu → feature)", "1-2 (minimal features)"],
    ["Feature badges", "NEW, HOT labels", "HOT, Promo labels", "NEW labels", "None"],
    ["Server-driven config", "Yes — recommended features personalized", "Yes — feature list server-controlled", "Yes — per region/mode", "No — static UI"],
]

table = doc.add_table(rows=len(matrix_data), cols=5)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, row_data in enumerate(matrix_data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(9)
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

# ============================================================
# 3. Per-Competitor Deep Dive
# ============================================================
doc.add_heading("3. Per-Competitor Deep Dive", level=1)

# --- Binance ---
doc.add_heading("3.1 Binance", level=2)

doc.add_heading("Navigation Structure", level=3)
doc.add_paragraph(
    "Binance uses a profile icon (top-left of homepage) that opens a full-screen Account Center. "
    "The Account Center serves as the primary hub for all non-trading features and settings."
)

doc.add_paragraph("Account Center layout (top to bottom):", style="List Bullet")
items = [
    "Profile section: identity verification status, Binance User ID (BUID), VIP status, security settings",
    "Customer support icon: 24/7 chat access",
    "Settings icon: app customization",
    "Shortcuts section: user-managed favorite mini apps with quick access from homepage",
    "Recommended features: personalized suggestions based on user behavior",
    "App mode toggle: switch between Pro and Lite interfaces",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet 2")

doc.add_paragraph(
    "Bottom tab bar: Home, Markets, Trade, Futures, Wallets (5 tabs). "
    "The homepage features quick action buttons for Buy, Trade, and Stake, plus live charts and a portfolio dashboard."
)

doc.add_heading("Notable UX Decisions", level=3)
doc.add_paragraph(
    "Binance separates 'Shortcuts' (user-curated) from 'Recommended' (system-suggested), giving users control "
    "while still surfacing feature discovery. The Pro/Lite toggle is prominently placed in the Account Center, "
    "acknowledging different user segments. Binance also launched AI-powered UI personalization in 2025, "
    "automatically adjusting the feature layout based on user behavior — an advanced approach that most "
    "competitors haven't matched."
)

doc.add_heading("Relevance to Aster", level=3)
doc.add_paragraph(
    "The Shortcuts + Recommended dual approach is worth considering for v2. For v1, Aster should focus on "
    "the user-managed shortcuts model. The Pro/Lite toggle is less relevant for Aster's DEX context where "
    "all users are assumed to be crypto-native."
)

# --- OKX ---
doc.add_heading("3.2 OKX", level=2)

doc.add_heading("Navigation Structure", level=3)
doc.add_paragraph(
    "OKX uses a profile avatar (top-left) that pushes a full-screen profile page onto the navigation stack. "
    "This is the model Danny specifically referenced as the target design."
)

doc.add_paragraph("Profile page layout (top to bottom):")
items = [
    "Profile header: avatar, username, verification badges (Verified, Regular user), QR code",
    "VIP promotion banner: \"Unlock OKX VIP\" with benefits (Discounted fees, Boosted yield, Priority support) — dismissible",
    "Shortcuts section: 4-icon grid with edit (pencil) button. Default: Get help, Demo trading, Affiliate, Campaign center. User-customizable via edit mode",
    "Manage assets section: 5-icon grid — Buy, Sell, Deposit, Withdraw, P2P trading",
    "Trade section: 8-icon grid in 2 rows — Convert, Spot, Futures, Options, Pre-market, Bots, Copy trading, Demo trading",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet 2")

doc.add_paragraph(
    "Bottom tab bar: 5 tabs (Home, Discover, Trade, Web3, Wallet). "
    "Settings are accessed via the profile page's \"Profile and settings\" link at the top."
)

doc.add_heading("Notable UX Decisions", level=3)
doc.add_paragraph(
    "OKX clearly separates feature categories (Shortcuts / Manage Assets / Trade) with horizontal dividers "
    "and section headers. The 'Shortcuts' section is the only user-customizable section — 'Manage Assets' and "
    "'Trade' sections have fixed layouts controlled by OKX. This is a smart balance: users personalize their "
    "most-used features while the core navigation remains consistent and predictable.\n\n"
    "The edit button (pencil icon) on Shortcuts triggers a drag-to-reorder mode. "
    "Each icon is 48x48dp with a text label below, arranged in a 4-column grid."
)

doc.add_heading("Relevance to Aster", level=3)
doc.add_paragraph(
    "This is the closest match to Aster's target design. Key adaptations needed:\n"
    "• Replace avatar + username with wallet address + identicon (DEX model)\n"
    "• Remove KYC/verification badges (not applicable to DEX)\n"
    "• Remove VIP banner (or replace with veASTER staking tier)\n"
    "• Add Aster-specific categories (Earn section for Shield/Staking, Explore for Referral/Leaderboard)\n"
    "• The 'Shortcuts' user-customization model should be adopted directly for v1"
)

# --- Bybit ---
doc.add_heading("3.3 Bybit", level=2)

doc.add_heading("Navigation Structure", level=3)
doc.add_paragraph(
    "Bybit offers dual modes: Lite (simplified for beginners) and Pro (full-featured for advanced traders). "
    "The mode toggle is at the top of the app, allowing instant switching."
)

doc.add_paragraph("Navigation differences by mode:")
items = [
    "Lite mode: clean homepage with one-click trading for major assets, customizable homepage widgets, step-by-step guidance",
    "Pro mode: full trading interface with Spot, Derivatives, TradFi access, integrated Copy trading, Bots, and ChatGPT strategy assistant. Customizable homepage layout",
    "My Page: profile, usage history, settings access — shared across both modes",
    "Earn section reorganized into 3 categories: Easy Earn, On-Chain Earn, Advanced Earn",
    "Exchange/Web3 toggle at top of main screen for switching between CEX and DeFi wallet",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet 2")

doc.add_paragraph(
    "Bottom tab bar: Home, Markets, Trade (or Derivatives), Copy, Assets (5 tabs)."
)

doc.add_heading("Notable UX Decisions", level=3)
doc.add_paragraph(
    "Bybit's Lite/Pro split is the most aggressive simplification approach among major CEXs. "
    "The customizable homepage in Pro mode lets users create their own widget layout, going beyond "
    "simple shortcut reordering. The Earn categorization (Easy/On-Chain/Advanced) is a strong pattern "
    "for Aster's Earn section which includes Shield, Earn, and Staking products of varying complexity."
)

doc.add_heading("Relevance to Aster", level=3)
doc.add_paragraph(
    "The Earn categorization pattern is worth adopting. The Lite/Pro split is over-engineered for "
    "Aster's DEX user base. The customizable homepage widgets concept could inform a v2 feature."
)

# --- Hyperliquid ---
doc.add_heading("3.4 Hyperliquid", level=2)

doc.add_heading("Navigation Structure", level=3)
doc.add_paragraph(
    "Hyperliquid is primarily a web-based DEX with no official native mobile app. "
    "The web interface is accessible via mobile browsers and third-party wrapper apps (OneShot, HyperTracker). "
    "Navigation is minimal and trading-focused."
)

doc.add_paragraph("Web interface navigation:")
items = [
    "Top nav bar: Trade, Portfolio, Vaults, Leaderboard, Points — flat horizontal navigation",
    "No hamburger menu or profile page — wallet connection is the only account concept",
    "Deposit/Transfer accessible within the trading interface",
    "No feature discovery grid or shortcuts — all features visible in the top nav",
    "No customization options for navigation",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet 2")

doc.add_heading("Notable UX Decisions", level=3)
doc.add_paragraph(
    "Hyperliquid prioritizes trading efficiency over feature breadth. The minimal navigation works because "
    "Hyperliquid offers far fewer features than a full-platform exchange. With only ~5 top-level sections, "
    "there's no need for a feature discovery menu. This is the opposite extreme from Binance/OKX."
)

doc.add_heading("Relevance to Aster", level=3)
doc.add_paragraph(
    "Hyperliquid demonstrates that DEX users expect a wallet-connected, minimal-chrome experience. "
    "Aster has more features than Hyperliquid, so a feature navigation menu is justified — but it should "
    "remain clean and avoid the information overload of CEX apps. Aster sits between Hyperliquid's "
    "minimalism and OKX's feature richness."
)

# ============================================================
# 4. Gap Analysis
# ============================================================
doc.add_heading("4. Gap Analysis", level=1)

doc.add_heading("What competitors offer that Aster doesn't (today)", level=3)
items = [
    "Feature navigation hub — all three CEXs have a dedicated feature discovery page; Aster's hamburger menu only has settings",
    "User-customizable shortcuts — Binance, OKX, and Bybit all allow users to personalize their quick-access features",
    "Categorized feature grid — competitors group features by type (Trade, Earn, Assets); Aster has no feature categorization",
    "Server-driven feature configuration — competitors can update feature lists without app releases",
    "Feature badges (NEW, HOT) — competitors highlight new or popular features to drive discovery",
    "WebView integration for web-only features — CEXs use in-app WebView for DeFi/Web3 features not built natively",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("What Aster offers that competitors don't", level=3)
items = [
    "DEX-native wallet experience — no separate KYC or account registration flow; wallet is the account",
    "Shield (principal-protected product) — unique product not offered by any competitor analyzed",
    "Simpler onboarding — one wallet address = one user, no account management complexity",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Where the market is converging", level=3)
items = [
    "Full-screen profile/account page replacing simple hamburger menus",
    "Feature grids organized by category (2-4 categories, 4-column icon grids)",
    "User customization of at least one section (shortcuts/favorites)",
    "Server-driven feature lists for OTA updates",
    "Dual-mode (simple/advanced) interfaces — though this is CEX-specific",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

# ============================================================
# 5. Recommendations for Aster
# ============================================================
doc.add_heading("5. Recommendations for Aster", level=1)

doc.add_heading("Model to follow: OKX profile page", level=3)
doc.add_paragraph(
    "OKX's profile page structure (Shortcuts → Manage Assets → Trade) is the best fit for Aster because:\n"
    "• Clear category separation with section headers\n"
    "• User-customizable shortcuts with a simple edit mode\n"
    "• Fixed categories for core features ensure discoverability\n"
    "• Full-screen push navigation (not drawer/overlay) provides ample space\n"
    "• Proven pattern that users of crypto apps already understand"
)

doc.add_heading("Key adaptations for Aster's DEX context", level=3)
items = [
    "Replace profile avatar/username with wallet address identicon — one wallet = one user (DEX model)",
    "Remove KYC badges — replace with veASTER staking tier badge if applicable",
    "Keep hamburger icon as trigger (per product decision) rather than OKX's avatar",
    "Add Aster-specific categories: Earn (Shield, Earn, Staking), Explore (Explorer, Referral, Rocket Launch, Leaderboard), Rewards (Trade & Earn, Points, Airdrop)",
    "Use WebView for features that exist on web but not yet in app — with auth token passthrough",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Suggested v1 scope based on competitive landscape", level=3)
items = [
    "Full-screen feature navigation page with categorized icon grid (must-have — all competitors have this)",
    "User-customizable shortcuts section with drag-to-reorder (must-have — Binance + OKX both support this)",
    "Server-driven feature list and ordering (must-have — enables OTA updates, all competitors do this)",
    "Feature badges (NEW, HOT) on icons (should-have — easy to implement, high impact on feature discovery)",
    "Settings sub-page relocated from hamburger menu (must-have — follows OKX pattern)",
    "WebView container for web-only features (must-have — bridges the app/web feature gap)",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Naming recommendations", level=3)
doc.add_paragraph(
    "Follow industry-standard naming for categories and features:\n"
    "• Use 'Deposit' / 'Withdraw' (not 'Fund' / 'Transfer') — matches Binance, OKX, Bybit\n"
    "• Use 'Perpetual' (not 'Futures' or 'Perp') — per Aster convention for full word\n"
    "• Use 'Spot' — universal across all competitors\n"
    "• Use 'Earn' as category name — matches Bybit and is intuitive\n"
    "• Use 'Explore' for discovery features — matches OKX's 'Discover' tab concept\n"
    "• 'Shield' is Aster-unique — no naming alignment needed, but should have a brief subtitle for new users"
)

doc.add_heading("Models to diverge from", level=3)
items = [
    "Binance's Pro/Lite mode split — over-engineered for Aster's crypto-native user base",
    "Binance's AI-powered personalization — v1 should use simple server-driven config, not ML",
    "Hyperliquid's minimal approach — Aster has too many features for a flat nav bar",
    "Bybit's homepage widget customization — too complex for v1; simple shortcut reorder is sufficient",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

# Save
output_path = "/Users/dannychen/Documents/aster-docs/docs/homepage-navigation-competitor-analysis.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
