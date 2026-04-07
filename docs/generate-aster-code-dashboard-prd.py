#!/usr/bin/env python3
"""Generate the Aster Code Dashboard PRD as .docx with embedded flow PNGs."""

from docx import Document
from docx.shared import Inches, Pt
import os

DOCS = "/Users/dannychen/Documents/aster-docs/docs"

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Helvetica Neue"
font.size = Pt(10)


def add_table(headers, rows):
    """Helper to add a styled table."""
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for run in t.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.rows[i + 1].cells[j].text = cell
    return t


def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


# ============================================================
# TITLE
# ============================================================
doc.add_heading("Aster Code Dashboard", level=0)

# ============================================================
# S1 Feature Overview
# ============================================================
doc.add_heading("1. Feature Overview", level=1)
doc.add_paragraph(
    "Build a public-facing Aster Code Dashboard that provides real-time analytics on Aster Code "
    "(referral code) performance across the Aster Exchange platform. The dashboard displays aggregate "
    "KPIs (total fees, volume, users), time-series charts broken down by individual Aster Code, "
    "proportion analysis relative to total exchange activity, and a sortable/paginated stats table "
    "listing all registered Aster Codes with their key metrics. This feature increases transparency "
    "for Aster Code partners, drives competition among referrers, and serves as a public signal of "
    "platform traction. Reference design: ASXN Hyperliquid Builder Codes dashboard."
)

doc.add_heading("Success Metrics", level=2)
metrics = [
    ["Dashboard daily unique visitors", ">500 within 30 days of launch", "Google Analytics / PostHog"],
    ["Aster Code partner retention (monthly active)", ">80% of registered partners visit dashboard at least once per month", "PostHog cohort analysis"],
    ["New Aster Code registrations", ">20% increase in new code registrations within 60 days of launch (dashboard as discovery tool)", "Internal registration data"],
    ["Referral volume growth", ">10% increase in total Aster Code volume within 90 days (competition effect)", "On-chain data / data pipeline"],
    ["Page load performance", "FCP <= 1.5s, full chart render <= 3s (p95)", "Lighthouse / RUM monitoring"],
]
add_table(["Metric", "Target", "Measurement"], metrics)

# ============================================================
# S2 Background & Motivation
# ============================================================
doc.add_heading("2. Background & Motivation", level=1)

doc.add_paragraph(
    "Aster Exchange supports an Aster Code (referral code) system where partners, integrators, "
    "and community members refer users to the platform and earn a share of trading fees generated "
    "by their referred users. Currently, Aster Code performance data is only accessible internally "
    "or through partner-specific dashboards with limited visibility."
)
doc.add_paragraph(
    "Hyperliquid's ASXN Builder Codes dashboard has set the industry benchmark for public referral "
    "analytics: it surfaces aggregate KPIs, per-builder breakdowns (volume, fees, users, trades), "
    "time-series trends, proportion analysis, and a comprehensive stats table with 386+ builder codes. "
    "This transparency creates a competitive ecosystem where builders actively promote and grow their "
    "referral networks."
)
doc.add_paragraph(
    "Aster needs a similar public dashboard to: (1) provide transparency and trust for existing "
    "Aster Code partners, (2) motivate partners to drive more volume through healthy competition, "
    "(3) serve as a public marketing tool showcasing platform growth and adoption, and (4) align "
    "with the DEX ethos of open, verifiable data."
)

# ============================================================
# S3 Scope
# ============================================================
doc.add_heading("3. Scope", level=1)

doc.add_heading("In Scope (v1)", level=2)
in_scope = [
    "Public dashboard page accessible without authentication at /aster-codes (or equivalent route)",
    "KPI summary cards: Total Aster Code Fees, Total Aster Code Volume, Total Aster Code Users",
    "6 time-series charts: Aster Code Volume (overall), Volume proportion, Volume by code, Fees by code, Cumulative users by code, User proportion",
    "2 code ranking bar charts: Volume by Code (descending bars + cumulative line), Fees by Code (descending bars + cumulative line)",
    "3 analysis charts: Average fees per user by code, Aster Code pie chart (switchable: fees/volume/trades/users), New users by code (bar ranking)",
    "All charts support time range switching (W / M / Q) and chart type toggling (line / bar / area where applicable)",
    "Legend-based code filtering: click code name to show/hide, 'Deselect all' button",
    "Cumulative volume overlay line on applicable bar charts",
    "Aster Code Stats table: sortable by all columns (Code, Address, Volume, Fees, Trade Count 30D, Total Users), paginated (10 per page)",
    "Copy-to-clipboard for wallet addresses in the stats table",
    "Dark theme UI consistent with Aster Exchange design system",
    "Responsive layout for desktop and mobile H5 browsers (breakpoints: desktop >= 1280px, tablet 768-1279px, mobile < 768px)",
    "Data pipeline: daily aggregation job that computes all dashboard metrics from on-chain / trading data",
]
add_bullets(in_scope)

doc.add_heading("Out of Scope (v2+)", level=2)
out_scope = [
    ("Individual Aster Code detail page", "v1 shows aggregate dashboard only; v2 adds drill-down page per code with historical performance"),
    ("Real-time streaming updates", "v1 uses daily batch data; v2 adds WebSocket streaming for near-real-time metrics"),
    ("Custom date range picker", "v1 supports preset ranges (W/M/Q); v2 adds custom start/end date selection"),
    ("Export to CSV/PDF", "v1 is view-only; v2 adds data export functionality"),
    ("Aster Code registration/management", "separate feature; this PRD covers the analytics dashboard only"),
    ("Authenticated partner dashboard", "v1 is fully public; v2 may add authenticated view with additional partner-specific metrics"),
    ("Comparison mode", "side-by-side comparison of two or more Aster Codes is v2"),
]
for title, reason in out_scope:
    doc.add_paragraph(f"{title} -- {reason}", style="List Bullet")

# ============================================================
# S4 User Stories
# ============================================================
doc.add_heading("4. User Stories", level=1)

stories = [
    ["Aster Code Partner", "Visit the Aster Code Dashboard", "See my code's ranking, volume, fees, and user count compared to other partners"],
    ["Aster Code Partner", "Switch the time range on charts to weekly", "Understand my code's short-term performance trend"],
    ["Aster Code Partner", "Click my code name in the chart legend", "Isolate my code's data from the stacked chart for clearer visibility"],
    ["Aster Code Partner", "Sort the stats table by fees descending", "Find the top-earning codes and benchmark my performance"],
    ["Aster Code Partner", "Copy a wallet address from the stats table", "Verify the on-chain address associated with a specific code"],
    ["Potential Partner", "Browse the dashboard without logging in", "Evaluate the Aster Code program's potential before applying"],
    ["Aster Team", "View the pie chart by volume dimension", "Understand volume concentration across codes for business planning"],
    ["Aster Team", "Check the proportion of exchange activity by codes", "Monitor how much of total exchange activity flows through Aster Codes"],
    ["Community Member", "Browse the dashboard", "See which Aster Codes are driving the most activity and growth on the platform"],
]
add_table(["Role", "Action", "Outcome"], stories)

# ============================================================
# S5 User Flows
# ============================================================
doc.add_heading("5. User Flows", level=1)

doc.add_heading("Flow 1: Browse Aster Code Dashboard", level=2)
doc.add_paragraph(
    "This flow covers the primary user journey: a visitor lands on the Aster Code Dashboard, "
    "sees a loading skeleton while data fetches, then browses KPI cards, scrolls through charts, "
    "and interacts with the stats table (sorting, paginating, and copying wallet addresses). "
    "Includes the empty state branch when no data is available."
)
flow1_path = os.path.join(DOCS, "flow-aster-code-dashboard-view.png")
if os.path.exists(flow1_path):
    doc.add_picture(flow1_path, width=Inches(5.5))

doc.add_heading("Flow 2: Filter & Interact with Charts", level=2)
doc.add_paragraph(
    "This flow covers chart interaction: switching time ranges (W/M/Q), toggling chart types "
    "(line/bar/area), switching pie chart dimensions (fees/volume/trades/users), and filtering "
    "codes via legend click. Includes the API retry flow on failure."
)
flow2_path = os.path.join(DOCS, "flow-aster-code-dashboard-filter.png")
if os.path.exists(flow2_path):
    doc.add_picture(flow2_path, width=Inches(5.5))

# ============================================================
# S6 Functional Requirements
# ============================================================
doc.add_heading("6. Functional Requirements", level=1)

# 6.1 KPI Cards
doc.add_heading("6.1 KPI Summary Cards", level=2)
add_bullets([
    "Display 3 KPI cards at the top of the page in a horizontal row.",
    "Card 1: Total Aster Code Fees -- aggregate fees generated across all Aster Codes, formatted with $ prefix and abbreviated (e.g., $69.57M).",
    "Card 2: Total Aster Code Volume -- aggregate trading volume routed through Aster Codes, formatted with $ prefix and abbreviated (e.g., $220.44B).",
    "Card 3: Total Aster Code Users -- total unique users referred via Aster Codes, abbreviated (e.g., 9.65M).",
    "Values must be fetched from the overview API endpoint and represent all-time totals.",
    "Number formatting: use K (thousands), M (millions), B (billions) abbreviations with 2 decimal places.",
])

# 6.2 Time-Series Charts
doc.add_heading("6.2 Time-Series Charts", level=2)
doc.add_paragraph("The dashboard must display 6 time-series charts arranged in a 2-column grid:")

chart_specs = [
    ["Aster Code Volume", "Bar chart", "Aggregate volume across all codes per time bucket", "W, M, Q"],
    ["Proportion of Aster Volume by Codes", "Area chart", "Percentage of total exchange volume flowing through Aster Codes", "W, M, Q"],
    ["Aster Code Volume by Code", "Stacked bar + cumulative line", "Volume breakdown by individual code with cumulative overlay", "W, M, Q, ALL"],
    ["Aster Code Fees by Code", "Stacked bar + cumulative line", "Fee breakdown by individual code with cumulative overlay", "W, M, Q, ALL"],
    ["Cumulative Users by Code", "Stacked bar", "Running total of users per code over time", "W, M, Q"],
    ["Proportion of Aster Users by Codes", "Bar chart", "Percentage of total exchange users referred through codes", "W, M, Q"],
]
add_table(["Chart Name", "Chart Type", "Description", "Time Ranges"], chart_specs)

doc.add_paragraph("")
add_bullets([
    "Each chart must have a chart type toggle (line / bar / area icons) in the top-left area of the chart card.",
    "Time range selector (W / M / Q / ALL where applicable) in the top-right area of the chart card.",
    "Default time range: M (monthly).",
    "Charts 3-6 must include a code legend showing all codes with color-coded indicators.",
    "Legend interaction: click a code name to toggle its visibility in the chart. Include a 'Deselect all' button.",
    "Charts 3-4 must show a cumulative volume/fee line overlaid on the stacked bars (right Y-axis).",
    "Loading skeleton must be shown while data is being fetched.",
])

# 6.3 Analysis Charts
doc.add_heading("6.3 Analysis Charts", level=2)
doc.add_paragraph("The dashboard must display 3 additional analysis charts:")

analysis_specs = [
    ["Average Fees per User by Code", "Multi-line chart", "Average fee per user over time, one line per code", "W, M, Q"],
    ["Aster Code Pie Chart", "Donut chart", "Distribution across codes; switchable dimension: Fees / Volume / Trades / Users", "D, W, M"],
    ["New Users by Code", "Horizontal bar ranking", "New user count per code for selected period, sorted descending", "W, M, Q"],
]
add_table(["Chart Name", "Chart Type", "Description", "Time Ranges"], analysis_specs)

doc.add_paragraph("")
add_bullets([
    "Pie chart must display percentage labels on each segment (e.g., 32%, 18%, 15%, 6%).",
    "Pie chart must have dimension tabs: Fees, Volume, Trades, Users. Default: Fees.",
    "Pie chart must support D (daily), W (weekly), M (monthly) time range.",
    "Average Fees per User chart must include code legend with toggle functionality.",
    "New Users bar chart must sort codes descending by new user count.",
])

# 6.4 Volume/Fees by Builder bar ranking
doc.add_heading("6.4 Code Ranking Bar Charts", level=2)
add_bullets([
    "Display 2 bar ranking charts: Aster Code Volume by Code (horizontal bars) and Aster Code Fees by Code (horizontal bars).",
    "Each chart shows all codes sorted descending by the respective metric.",
    "Include a cumulative line overlay (right Y-axis) showing cumulative sum across codes.",
    "Support D (daily), W (weekly), M (monthly) time range switching.",
    "Toggle between Volume and Cumulative Volume (or Fees and Cumulative Fees).",
])

# 6.5 Stats Table
doc.add_heading("6.5 Aster Code Stats Table", level=2)
add_bullets([
    "Display a full-width data table titled 'Aster Code Stats' below all charts.",
    "Table columns:",
])

col_specs = [
    ["ASTER CODE", "String", "Display name of the Aster Code", "Alphabetical ascending"],
    ["ADDRESS", "String (truncated)", "Wallet address associated with the code (e.g., 0x1924...80e5)", "Alphabetical ascending"],
    ["VOLUME", "Currency ($)", "All-time trading volume routed through this code", "Numeric descending (default sort)"],
    ["FEES", "Currency ($)", "All-time fees generated by this code", "Numeric descending"],
    ["TRADE COUNT (30D)", "Number", "Number of trades in the last 30 days", "Numeric descending"],
    ["TOTAL USERS", "Number", "Total unique users referred by this code", "Numeric descending"],
]
add_table(["Column", "Type", "Description", "Sort Order"], col_specs)

doc.add_paragraph("")
add_bullets([
    "Default sort: Volume descending.",
    "All column headers must be clickable to sort ascending/descending. Show sort indicator (arrow) on active column.",
    "Pagination: 10 rows per page. Show 'Showing X-Y of Z' label and page navigation (previous/next + page number).",
    "Address column must have a copy icon button. On click, copy the full (non-truncated) wallet address to clipboard and show a brief toast: 'Address copied'.",
    "Table must handle large datasets (potentially hundreds of codes) with server-side pagination.",
])

# 6.6 Mobile H5 Responsive Layout
doc.add_heading("6.6 Mobile H5 Responsive Layout", level=2)
add_bullets([
    "The dashboard must be fully functional on mobile H5 browsers (Safari, Chrome mobile) with responsive breakpoints: desktop >= 1280px, tablet 768-1279px, mobile < 768px.",
    "KPI cards: stack vertically on mobile (1 column), side-by-side on tablet (3 columns).",
    "Chart grid: switch from 2-column to 1-column layout on mobile. Charts must maintain minimum height of 280px on mobile for readability.",
    "Time range and chart type toggles: remain accessible on mobile. Use compact toggle buttons (not dropdowns).",
    "Chart legends: collapsible on mobile to save vertical space. Default collapsed, tap to expand.",
    "Pie chart: reduce label density on mobile; show top 5 segments with labels, rest grouped as 'Others'.",
    "Stats table: horizontally scrollable on mobile. Code name and Volume columns remain sticky (fixed left) during horizontal scroll.",
    "Table pagination: use simplified pagination on mobile (previous / next only, no page numbers).",
    "Copy address button: increase tap target to minimum 44x44dp on mobile for accessibility.",
    "Touch interactions: support swipe gestures for chart scrolling on mobile. Disable hover-dependent interactions (tooltips trigger on tap instead).",
])

# 6.7 Data Pipeline
doc.add_heading("6.7 Data Pipeline", level=2)
add_bullets([
    "A daily aggregation job must run to compute all dashboard metrics from on-chain trading data.",
    "Aggregation must compute: per-code daily volume, fees, trade count, new users, and cumulative totals.",
    "Aggregated data must be stored in a dedicated analytics table/store optimized for read-heavy dashboard queries.",
    "Proportion metrics must be computed relative to total exchange volume/users for the same period.",
    "The pipeline must be idempotent: re-running for the same date produces identical results.",
    "Data freshness: dashboard data must reflect up to the previous day's complete trading activity (T-1).",
    "Pipeline failure handling: if the daily job fails, the dashboard shows the last successful data set. Alert the ops team via monitoring (see S7.6 pipeline-status endpoint).",
    "Code visibility filtering: each Aster Code record must have an 'isPublic' boolean flag. The pipeline must exclude codes where isPublic=false from all dashboard aggregations. Default: isPublic=true for all codes. Internal/test codes can be hidden by setting isPublic=false in the code registry.",
])

# ============================================================
# S7 API Specification
# ============================================================
doc.add_heading("7. API Specification", level=1)

doc.add_paragraph(
    "All endpoints follow the standard Aster API envelope pattern. These are public endpoints "
    "requiring no authentication."
)
doc.add_paragraph("Response envelope:")
doc.add_paragraph('{ "code": 0, "data": { ... }, "message": "success" }')

# --- 7.1 GET /api/v1/aster-code/overview ---
doc.add_heading("7.1 GET /api/v1/aster-code/overview", level=2)
doc.add_paragraph("Returns aggregate KPI totals for all Aster Codes.")
doc.add_paragraph("Request: No parameters.")
doc.add_paragraph("Response data fields:")
add_table(
    ["Field", "Type", "Description"],
    [
        ["totalFees", "string", "Total fees in USD, e.g., '69570000.00'"],
        ["totalVolume", "string", "Total volume in USD, e.g., '220440000000.00'"],
        ["totalUsers", "integer", "Total unique referred users, e.g., 9650000"],
        ["lastUpdatedAt", "string", "ISO 8601 timestamp of last successful pipeline run"],
    ],
)

# --- 7.2 GET /api/v1/aster-code/series (consolidated) ---
doc.add_heading("7.2 GET /api/v1/aster-code/series", level=2)
doc.add_paragraph(
    "Unified time-series endpoint for volume, fees, users, and proportion data. "
    "Consolidates what would otherwise be 4 separate endpoints into a single flexible API."
)

doc.add_paragraph("Request query parameters:")
add_table(
    ["Param", "Type", "Required", "Description"],
    [
        ["metric", "enum", "Yes", "'volume', 'fees', 'users', or 'proportion'"],
        ["interval", "enum", "No", "Time bucket: 'W' (week), 'M' (month, default), 'Q' (quarter)"],
        ["breakdown", "boolean", "No", "If true, returns per-code breakdown. Default: false (aggregate)"],
        ["subMetric", "enum", "No", "For metric=users: 'cumulative' (default) or 'new'. For metric=proportion: 'volume' (default) or 'users'. Ignored for other metrics."],
    ],
)

doc.add_paragraph("Response data fields (aggregate mode, breakdown=false):")
add_table(
    ["Field", "Type", "Description"],
    [
        ["series", "array", "Array of time-bucketed data points"],
        ["series[].date", "string", "Bucket start date (ISO 8601, e.g., '2026-03-01')"],
        ["series[].value", "string", "Metric value for this bucket (volume in USD, fees in USD, user count, or proportion as decimal)"],
    ],
)

doc.add_paragraph("Response data fields (breakdown mode, breakdown=true):")
add_table(
    ["Field", "Type", "Description"],
    [
        ["series", "array", "Array of time-bucketed data points"],
        ["series[].date", "string", "Bucket start date"],
        ["series[].codes", "array", "Per-code metric values"],
        ["series[].codes[].code", "string", "Aster Code name"],
        ["series[].codes[].value", "string", "Metric value for this code in this bucket"],
        ["series[].cumulativeValue", "string", "Running cumulative total across all codes (when applicable)"],
    ],
)

doc.add_paragraph("Response data fields (proportion mode, metric=proportion):")
add_table(
    ["Field", "Type", "Description"],
    [
        ["series", "array", "Array of time-bucketed data points"],
        ["series[].date", "string", "Bucket start date"],
        ["series[].proportion", "string", "Proportion as decimal (e.g., '0.0523' for 5.23%)"],
        ["series[].total", "string", "Total exchange metric value for context"],
        ["series[].asterCodeTotal", "string", "Aster Code metric value for context"],
    ],
)

# --- 7.3 GET /api/v1/aster-code/ranking ---
doc.add_heading("7.3 GET /api/v1/aster-code/ranking", level=2)
doc.add_paragraph("Returns per-code ranking data for bar charts and pie chart.")

doc.add_paragraph("Request query parameters:")
add_table(
    ["Param", "Type", "Required", "Description"],
    [
        ["metric", "enum", "Yes", "'volume', 'fees', 'trades', or 'users'"],
        ["interval", "enum", "No", "'D' (daily), 'W' (weekly), 'M' (monthly, default)"],
        ["limit", "integer", "No", "Max codes to return (default: 20, max: 50)"],
    ],
)

doc.add_paragraph("Response data fields:")
add_table(
    ["Field", "Type", "Description"],
    [
        ["rankings", "array", "Array of codes sorted descending by metric"],
        ["rankings[].code", "string", "Aster Code name"],
        ["rankings[].value", "string", "Metric value for the selected period"],
        ["rankings[].cumulativeValue", "string", "Cumulative sum up to and including this code"],
        ["rankings[].proportion", "string", "This code's share of total (decimal)"],
    ],
)

# --- 7.4 GET /api/v1/aster-code/avg-fee-per-user ---
doc.add_heading("7.4 GET /api/v1/aster-code/avg-fee-per-user", level=2)
doc.add_paragraph("Returns time-series average fee per user for each code.")

doc.add_paragraph("Request query parameters:")
add_table(
    ["Param", "Type", "Required", "Description"],
    [
        ["interval", "enum", "No", "'W', 'M' (default), 'Q'"],
    ],
)

doc.add_paragraph("Response data fields:")
add_table(
    ["Field", "Type", "Description"],
    [
        ["series", "array", "Array of time-bucketed data points"],
        ["series[].date", "string", "Bucket start date"],
        ["series[].codes", "array", "Per-code average fee"],
        ["series[].codes[].code", "string", "Aster Code name"],
        ["series[].codes[].avgFeePerUser", "string", "Average fee per user in USD"],
    ],
)

# --- 7.5 GET /api/v1/aster-code/stats ---
doc.add_heading("7.5 GET /api/v1/aster-code/stats", level=2)
doc.add_paragraph("Returns paginated table data for the Aster Code Stats section.")

doc.add_paragraph("Request query parameters:")
add_table(
    ["Param", "Type", "Required", "Description"],
    [
        ["page", "integer", "No", "Page number (1-based, default: 1)"],
        ["pageSize", "integer", "No", "Items per page (default: 10, max: 50)"],
        ["sortBy", "enum", "No", "'volume' (default), 'fees', 'tradeCount30d', 'totalUsers', 'code'"],
        ["sortOrder", "enum", "No", "'desc' (default) or 'asc'"],
    ],
)

doc.add_paragraph("Response data fields:")
add_table(
    ["Field", "Type", "Description"],
    [
        ["items", "array", "Array of code stats for current page"],
        ["items[].code", "string", "Aster Code display name"],
        ["items[].address", "string", "Full wallet address (not truncated)"],
        ["items[].volume", "string", "All-time volume in USD"],
        ["items[].fees", "string", "All-time fees in USD"],
        ["items[].tradeCount30d", "integer", "Trade count in last 30 days"],
        ["items[].totalUsers", "integer", "Total referred users"],
        ["total", "integer", "Total number of Aster Codes"],
        ["page", "integer", "Current page number"],
        ["pageSize", "integer", "Items per page"],
        ["totalPages", "integer", "Total number of pages"],
    ],
)

# --- 7.6 GET /api/v1/aster-code/pipeline-status (internal) ---
doc.add_heading("7.6 GET /api/v1/aster-code/pipeline-status", level=2)
doc.add_paragraph(
    "Internal endpoint for ops/admin monitoring of the data pipeline health. "
    "Requires admin authentication (10xxx auth). Not exposed to the public dashboard."
)

doc.add_paragraph("Request: Requires admin JWT token in Authorization header.")
doc.add_paragraph("Response data fields:")
add_table(
    ["Field", "Type", "Description"],
    [
        ["status", "enum", "'healthy', 'degraded', 'failed'"],
        ["lastSuccessfulRun", "string", "ISO 8601 timestamp of last successful pipeline completion"],
        ["lastAttemptedRun", "string", "ISO 8601 timestamp of last pipeline attempt"],
        ["nextScheduledRun", "string", "ISO 8601 timestamp of next scheduled run"],
        ["recordCount", "integer", "Total Aster Codes in the aggregated dataset"],
        ["errorMessage", "string | null", "Error message if last run failed, null otherwise"],
    ],
)
doc.add_paragraph(
    "Monitoring: Ops team should set up alerting on this endpoint. If status != 'healthy' "
    "for more than 2 consecutive checks (i.e., pipeline down >48h), escalate to engineering."
)

# ============================================================
# S8 Error Handling
# ============================================================
doc.add_heading("8. Error Handling", level=1)

doc.add_paragraph(
    "All API errors follow the standard envelope with non-zero code. Dashboard/analytics error "
    "codes use the 60xxx domain (new domain: 60xxx = Analytics & Dashboard), consistent with "
    "the existing 5-digit error code convention (10xxx auth, 20xxx trading, 30xxx staking, 40xxx wallet)."
)

err_data = [
    ["60001", "Dashboard data unavailable", "数据暂不可用", "資料暫不可用", "Data pipeline has not run or failed; no data available"],
    ["60002", "Invalid time range", "无效的时间范围", "無效的時間範圍", "Client sends unsupported interval parameter"],
    ["60003", "Invalid sort parameter", "无效的排序参数", "無效的排序參數", "Client sends unsupported sortBy or sortOrder value"],
    ["60004", "Page out of range", "页码超出范围", "頁碼超出範圍", "Requested page exceeds total pages"],
    ["60005", "Invalid metric parameter", "无效的指标参数", "無效的指標參數", "Client sends unsupported metric value"],
]
add_table(["Code", "EN", "zh-Hans", "zh-Hant", "Trigger"], err_data)

doc.add_paragraph("")
doc.add_heading("Edge Cases", level=2)
edges = [
    "No data available (new deployment, pipeline not yet run): KPI cards show $0 / 0. Charts show empty state with message: 'No data yet. Check back after the first data sync.'",
    "Pipeline partially failed (some metrics available, others not): show available data, display inline error for failed chart sections with a 'Retry' button.",
    "Extremely large number of Aster Codes (1000+): stats table uses server-side pagination to avoid performance issues. Charts show top N codes with an 'Others' aggregation bucket.",
    "Single code dominates all metrics (>90% share): pie chart still renders all segments; no special treatment. The proportion is clearly visible from the percentage labels.",
    "Aster Code with zero activity: still appears in stats table with 0 values. Does not appear in ranking charts (filtered out of top N).",
    "Stale data (pipeline down for >24h): dashboard shows the last successful data set. A subtle banner at the top indicates: 'Data last updated: [timestamp]. Updates may be delayed.'",
]
add_bullets(edges)

# ============================================================
# S9 Non-Functional Requirements
# ============================================================
doc.add_heading("9. Non-Functional Requirements", level=1)

nfr_data = [
    ["Performance -- initial page load", "Time to First Contentful Paint (FCP) <= 1.5s. KPI cards render within 2s. All charts render within 3s on standard broadband."],
    ["Performance -- chart interaction", "Time range switching and chart type toggling must complete within 500ms (API response + re-render)."],
    ["Performance -- table pagination", "Page navigation must complete within 300ms."],
    ["Data freshness", "Dashboard data reflects T-1 (previous day's complete trading data). Updated once daily via batch pipeline."],
    ["Scalability", "API must handle 100+ concurrent dashboard viewers without degradation. Use caching (CDN or Redis) for all public read endpoints."],
    ["Caching", "API responses cached with 1-hour TTL (data updates daily, so aggressive caching is safe). Cache invalidated when daily pipeline completes."],
    ["i18n", "All user-facing strings (chart titles, labels, error messages, table headers) must support EN, zh-Hans, zh-Hant. No hardcoded user-facing strings."],
    ["SEO", "Dashboard page must be server-side rendered (SSR) or statically generated (SSG) for search engine indexing. Page title, description, and Open Graph meta tags must be set."],
    ["Accessibility", "Charts must have aria-labels describing their content. Table must be navigable via keyboard. Color-blind-friendly palette for chart series."],
    ["Security", "All API endpoints are public (no auth required). Rate limit: 60 requests/minute per IP to prevent abuse. No PII exposed (wallet addresses are public on-chain data)."],
    ["Browser support", "Desktop: Chrome, Firefox, Safari, Edge (latest 2 versions). Mobile H5: Safari iOS, Chrome Android (latest 2 versions). Minimum viewport width: 320px (iPhone SE)."],
    ["Chart library", "Use a production-grade charting library (e.g., Recharts, Apache ECharts, or Highcharts) that supports all required chart types, legend interaction, and responsive sizing."],
]
add_table(["Dimension", "Requirement"], nfr_data)

# ============================================================
# S10 Open Questions
# ============================================================
doc.add_heading("10. Open Questions", level=1)

questions = [
    ["1", "What is the exact URL path for the dashboard? /aster-codes, /dashboard/aster-codes, or /analytics/aster-codes?", "Routing configuration, SEO", "/aster-codes (short, clean)"],
    ["2", "How is an Aster Code registered? Is there an existing registration flow, or does the dashboard only display codes that have been registered through a separate process?", "Determines whether the dashboard needs a CTA or link to registration", "Display-only in v1; link to registration docs if available"],
    ["3", "What charting library does the Next.js admin dashboard currently use? Should the public dashboard use the same library for consistency?", "Development effort, visual consistency", "Use the same library as the admin dashboard if suitable; otherwise choose ECharts for its rich feature set"],
    ["4", "Should the chart color palette for individual codes be deterministic (same code always gets same color) or dynamic?", "Visual consistency across visits", "Deterministic: hash the code name to assign a stable color from the palette"],
    ["5", "What is the expected number of Aster Codes at launch? This affects whether client-side or server-side pagination/aggregation is needed for charts.", "Architecture decision", "Assume 100+ codes at launch; use server-side pagination for table, top-N (N=15) + Others for charts"],
    ["6", "Should the 'Proportion' charts show proportion of total exchange activity or only perpetual futures volume?", "Metric definition", "Total exchange activity (all products: perpetual + spot + Shield)"],
]
add_table(["#", "Question", "Impact", "Proposed Default"], questions)

# ============================================================
# Save
# ============================================================
output_path = os.path.join(DOCS, "aster-code-dashboard-prd.docx")
doc.save(output_path)
print(f"Saved to {output_path}")
